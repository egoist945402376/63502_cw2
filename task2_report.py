from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)

def add_heading_black(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLACK
    return p

def add_code_block(doc, code):
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return para

def add_intro(doc, text):
    p = doc.add_paragraph()
    r1 = p.add_run('Introduction: ')
    r1.bold = True
    r1.font.color.rgb = BLACK
    r2 = p.add_run(text)
    r2.font.color.rgb = BLACK
    return p

# ── document ──────────────────────────────────────────────────────────────────
doc = Document()

add_heading_black(doc, 'Task 2 (City)', level=3)

# ── Task 2.1 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 2.1', level=4)

add_intro(doc,
    'Queries dbo:City instances linked to the UK via dbo:country, '
    'using FILTER IN to cover England, Scotland, Wales and Northern Ireland, '
    'with English labels only.'
)

add_code_block(doc, """\
PREFIX dbo:  <http://dbpedia.org/ontology/>
PREFIX dbr:  <http://dbpedia.org/resource/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?city ?cityName
WHERE {
  ?city a dbo:City ;
        dbo:country ?country ;
        rdfs:label ?cityName .
  FILTER (?country IN (dbr:United_Kingdom, dbr:England, dbr:Scotland,
                       dbr:Wales, dbr:Northern_Ireland))
  FILTER (LANG(?cityName) = "en")
}
ORDER BY ?cityName""")

# ── Task 2.2 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 2.2', level=4)

add_intro(doc,
    'Extends Task 2.1 with OPTIONAL clauses for population (dbo:populationTotal), '
    'area (dbo:areaTotal), and geographic coordinates (geo:lat / geo:long).'
)

add_code_block(doc, """\
PREFIX dbo:  <http://dbpedia.org/ontology/>
PREFIX dbr:  <http://dbpedia.org/resource/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
PREFIX geo:  <http://www.w3.org/2003/01/geo/wgs84_pos#>

SELECT DISTINCT ?city ?cityName ?population ?area ?lat ?long
WHERE {
  ?city a dbo:City ;
        dbo:country ?country ;
        rdfs:label ?cityName .
  FILTER (?country IN (dbr:United_Kingdom, dbr:England, dbr:Scotland,
                       dbr:Wales, dbr:Northern_Ireland))
  FILTER (LANG(?cityName) = "en")
  OPTIONAL { ?city dbo:populationTotal ?population }
  OPTIONAL { ?city dbo:areaTotal       ?area }
  OPTIONAL { ?city geo:lat             ?lat }
  OPTIONAL { ?city geo:long            ?long }
}
ORDER BY ?cityName""")

# ── Task 2.3 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 2.3', level=4)

p = doc.add_paragraph()
r = p.add_run('Third query (Wikidata city list, equivalent of Task 2.1): ')
r.bold = True
r.font.color.rgb = BLACK
p.add_run(
    'Retrieves instances of city (Q515) or any subclass via P31/P279*, '
    'filtered to country = United Kingdom (Q145), with English labels.'
).font.color.rgb = BLACK

add_code_block(doc, """\
PREFIX wd:   <http://www.wikidata.org/entity/>
PREFIX wdt:  <http://www.wikidata.org/prop/direct/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?city ?cityName
WHERE {
  ?city wdt:P31/wdt:P279* wd:Q515 ;
        wdt:P17 wd:Q145 ;
        rdfs:label ?cityName .
  FILTER (LANG(?cityName) = "en")
}
ORDER BY ?cityName""")

p = doc.add_paragraph()
r = p.add_run('Fourth query (Wikidata city attributes, equivalent of Task 2.2): ')
r.bold = True
r.font.color.rgb = BLACK
p.add_run(
    'Extends the third query with population (P1082), area (P2046) '
    'and geographic coordinates (P625).'
).font.color.rgb = BLACK

add_code_block(doc, """\
PREFIX wd:   <http://www.wikidata.org/entity/>
PREFIX wdt:  <http://www.wikidata.org/prop/direct/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?city ?cityName ?population ?area ?coord
WHERE {
  ?city wdt:P31/wdt:P279* wd:Q515 ;
        wdt:P17 wd:Q145 ;
        rdfs:label ?cityName .
  FILTER (LANG(?cityName) = "en")
  OPTIONAL { ?city wdt:P1082 ?population }
  OPTIONAL { ?city wdt:P2046 ?area }
  OPTIONAL { ?city wdt:P625  ?coord }
}
ORDER BY ?cityName""")

# ── save ──────────────────────────────────────────────────────────────────────
out_path = r'C:\Users\94540\Desktop\UoM\63502\cw2\task2.docx'
doc.save(out_path)
print(f'Saved to {out_path}')
