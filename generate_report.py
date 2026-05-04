from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)

def add_heading_black(doc, text, level):
    """Add a heading with forced black colour."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLACK
    return p

def add_code_block(doc, code):
    """Add a grey-background monospace paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return para

def add_intro(doc, text):
    """Add an 'Introduction:' label followed by body text."""
    p = doc.add_paragraph()
    r1 = p.add_run('Introduction: ')
    r1.bold = True
    r1.font.color.rgb = BLACK
    r2 = p.add_run(text)
    r2.font.color.rgb = BLACK
    return p

def _cell_bold(cell, text):
    """Write bold black text into a table cell (clears existing content)."""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.color.rgb = BLACK

def _cell_text(cell, text):
    """Write normal black text into a table cell (clears existing content)."""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.color.rgb = BLACK

# ── document ──────────────────────────────────────────────────────────────────
doc = Document()

# ---------- Title ----------
title = doc.add_heading('CW2 Report: Knowledge Graph and Semantic Table Annotation', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = BLACK

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# Part I
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_black(doc, 'Part I: SPARQL Query', level=2)
add_heading_black(doc, 'Task 1 (Country)', level=3)

# ── Task 1.1 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 1.1', level=4)

add_intro(doc,
    'This query selects instances of dbo:Country that have a capital city, '
    'excludes historical countries by filtering out those with a dissolutionYear, '
    'and restricts labels to English.'
)

add_code_block(doc, """\
PREFIX dbo:  <http://dbpedia.org/ontology/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?country ?countryName
WHERE {
  ?country a dbo:Country ;
           rdfs:label ?countryName ;
           dbo:capital ?capital .
  FILTER NOT EXISTS { ?country dbo:dissolutionYear ?year }
  FILTER (LANG(?countryName) = "en")
}
ORDER BY ?countryName""")

p = doc.add_paragraph()
r = p.add_run('Countries #: ')
r.bold = True
r.font.color.rgb = BLACK
rn = p.add_run('~195 (varies by DBpedia snapshot)')
rn.font.color.rgb = BLACK

# ── Task 1.2 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 1.2', level=4)

add_intro(doc,
    'Built on the Task 1.1 filter, this query uses OPTIONAL clauses to retrieve '
    'five country attributes: name (rdfs:label), ISO code (dbo:iso31661Code), '
    'telephone calling code (dbp:callingCode), area in km² (dbp:areaKm), '
    'and population estimate (dbp:populationEstimate).'
)

add_code_block(doc, """\
PREFIX dbo:  <http://dbpedia.org/ontology/>
PREFIX dbp:  <http://dbpedia.org/property/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?country ?countryName ?isoCode ?callingCode ?area ?population
WHERE {
  ?country a dbo:Country ;
           rdfs:label ?countryName ;
           dbo:capital ?capital .
  FILTER NOT EXISTS { ?country dbo:dissolutionYear ?year }
  FILTER (LANG(?countryName) = "en")
  OPTIONAL { ?country dbo:iso31661Code       ?isoCode }
  OPTIONAL { ?country dbp:callingCode        ?callingCode }
  OPTIONAL { ?country dbp:areaKm             ?area }
  OPTIONAL { ?country dbp:populationEstimate ?population }
}
ORDER BY ?countryName""")

# ── Task 1.3 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 1.3', level=4)

# Third query – Wikidata country list (equivalent of Task 1.1)
p = doc.add_paragraph()
r = p.add_run('Third query (Wikidata country list, equivalent of Task 1.1): ')
r.bold = True
r.font.color.rgb = BLACK
r2 = p.add_run(
    'Filters sovereign states (wd:Q3624078) that have a capital (P36), '
    'excludes dissolved entities (P576 absent), and restricts labels to English.'
)
r2.font.color.rgb = BLACK

add_code_block(doc, """\
PREFIX wd:   <http://www.wikidata.org/entity/>
PREFIX wdt:  <http://www.wikidata.org/prop/direct/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?country ?countryName
WHERE {
  ?country wdt:P31 wd:Q3624078 ;
           wdt:P36 ?capital ;
           rdfs:label ?countryName .
  FILTER NOT EXISTS { ?country wdt:P576 ?dissolved }
  FILTER (LANG(?countryName) = "en")
}
ORDER BY ?countryName""")

# Fourth query – Wikidata country attributes + provinces (equivalent of Task 1.2)
p = doc.add_paragraph()
r = p.add_run('Fourth query (Wikidata country attributes + provinces, equivalent of Task 1.2): ')
r.bold = True
r.font.color.rgb = BLACK
r2 = p.add_run(
    'Extends the third query with OPTIONAL clauses for ISO-3166-1 alpha-2 code (P297), '
    'calling code (P474), area (P2046) and population (P1082). '
    'Province names are gathered via P150 (contains administrative territorial entity) '
    'and concatenated with GROUP_CONCAT.'
)
r2.font.color.rgb = BLACK

add_code_block(doc, """\
PREFIX wd:   <http://www.wikidata.org/entity/>
PREFIX wdt:  <http://www.wikidata.org/prop/direct/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?country ?countryName ?isoCode ?callingCode ?area ?population
       (GROUP_CONCAT(DISTINCT ?provinceName; SEPARATOR="; ") AS ?provinces)
WHERE {
  ?country wdt:P31 wd:Q3624078 ;
           wdt:P36 ?capital ;
           rdfs:label ?countryName .
  FILTER NOT EXISTS { ?country wdt:P576 ?dissolved }
  FILTER (LANG(?countryName) = "en")
  OPTIONAL { ?country wdt:P297  ?isoCode }
  OPTIONAL { ?country wdt:P474  ?callingCode }
  OPTIONAL { ?country wdt:P2046 ?area }
  OPTIONAL { ?country wdt:P1082 ?population }
  OPTIONAL {
    ?country wdt:P150 ?province .
    ?province rdfs:label ?provinceName .
    FILTER (LANG(?provinceName) = "en")
  }
}
GROUP BY ?country ?countryName ?isoCode ?callingCode ?area ?population
ORDER BY ?countryName""")

# ── Task 1.4 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 1.4', level=4)

doc.add_paragraph().add_run(
    'The table below summarises approximate time spent and key observations for each of the four SPARQL queries.'
).font.color.rgb = BLACK

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

# Header row
hdr = table.rows[0].cells
_cell_bold(hdr[0], 'SPARQL Query')
_cell_bold(hdr[1], 'Time Spent')
_cell_bold(hdr[2], 'Comments')

# Data rows
rows_data = [
    (
        'Task 1.1\n(First DBpedia query to get countries)',
        '~1.5 hours',
        'The most time-consuming job is identifying the right filter combination '
        '(requiring dbo:capital and excluding dbo:dissolutionYear), because many '
        'dbo:Country instances are historical entities. '
        'Challenge: DBpedia does not have a single reliable "sovereign state" class, '
        'so proxy filters must be used instead.'
    ),
    (
        'Task 1.2\n(Second DBpedia query for country attributes)',
        '~2 hours',
        'The most time-consuming job is locating the correct property names for ISO '
        'code and calling code; DBpedia mixes dbo: and dbp: namespaces '
        '(e.g. dbp:isoCode vs dbo:isoCode) with inconsistent coverage. '
        'Lesson: always check both namespaces and use OPTIONAL to avoid losing rows.'
    ),
    (
        'Task 1.3 – third query\n(Wikidata country list)',
        '~1 hour',
        'The most time-consuming job is learning Wikidata Q-codes: '
        'sovereign state = Q3624078, capital = P36, dissolution date = P576. '
        'Challenge: Wikidata uses numeric property/entity IDs with no readable '
        'names in the query, making it harder to read and debug than DBpedia.'
    ),
    (
        'Task 1.3 – fourth query\n(Wikidata attributes + provinces)',
        '~2.5 hours',
        'The most time-consuming job is implementing GROUP_CONCAT for province names '
        'while correctly listing all non-aggregated variables in GROUP BY. '
        'Challenge: P150 (contains administrative territorial entity) includes '
        'many subdivision levels, so results may contain regions, districts, '
        'and municipalities, not just top-level provinces.'
    ),
]

for i, (query, time, comment) in enumerate(rows_data, start=1):
    row = table.rows[i].cells
    _cell_text(row[0], query)
    _cell_text(row[1], time)
    _cell_text(row[2], comment)

# ═══════════════════════════════════════════════════════════════════════════════
# Task 2 (City)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_black(doc, 'Task 2 (City)', level=3)

# ── Task 2.1 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 2.1', level=4)

add_intro(doc,
    'Queries dbo:City instances linked to the UK via dbo:country, '
    'using a FILTER IN to cover England, Scotland, Wales and Northern Ireland, '
    'with English labels only.'
)

add_code_block(doc, """\
PREFIX dbo:  <http://dbpedia.org/ontology/>
PREFIX dbr:  <http://dbpedia.org/resource/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?city ?cityName
WHERE {
  ?city a dbo:City ;
        dbo:country ?country ;
        rdfs:label ?cityName .
  FILTER (?country IN (dbr:United_Kingdom, dbr:England, dbr:Scotland,
                       dbr:Wales, dbr:Northern_Ireland))
  FILTER (LANG(?cityName) = "en")
}
ORDER BY ?cityName""")

# ── Task 2.2 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 2.2', level=4)

add_intro(doc,
    'Extends Task 2.1 with OPTIONAL clauses for population (dbo:populationTotal), '
    'area (dbo:areaTotal), and coordinates (geo:lat / geo:long).'
)

add_code_block(doc, """\
PREFIX dbo:  <http://dbpedia.org/ontology/>
PREFIX dbr:  <http://dbpedia.org/resource/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>
PREFIX geo:  <http://www.w3.org/2003/01/geo/wgs84_pos#>

SELECT DISTINCT ?city ?cityName ?population ?area ?lat ?long
WHERE {
  ?city a dbo:City ;
        dbo:country ?country ;
        rdfs:label ?cityName .
  FILTER (?country IN (dbr:United_Kingdom, dbr:England, dbr:Scotland,
                       dbr:Wales, dbr:Northern_Ireland))
  FILTER (LANG(?cityName) = "en")
  OPTIONAL { ?city dbo:populationTotal ?population }
  OPTIONAL { ?city dbo:areaTotal       ?area }
  OPTIONAL { ?city geo:lat             ?lat }
  OPTIONAL { ?city geo:long            ?long }
}
ORDER BY ?cityName""")

# ── Task 2.3 ──────────────────────────────────────────────────────────────────
add_heading_black(doc, 'Task 2.3', level=4)

p = doc.add_paragraph()
r = p.add_run('Third query (Wikidata city list, equivalent of Task 2.1): ')
r.bold = True
r.font.color.rgb = BLACK
r2 = p.add_run(
    'Retrieves instances of city (Q515) or any subclass via P31/P279*, '
    'filtered to country = United Kingdom (Q145), with English labels.'
)
r2.font.color.rgb = BLACK

add_code_block(doc, """\
PREFIX wd:   <http://www.wikidata.org/entity/>
PREFIX wdt:  <http://www.wikidata.org/prop/direct/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?city ?cityName
WHERE {
  ?city wdt:P31/wdt:P279* wd:Q515 ;
        wdt:P17 wd:Q145 ;
        rdfs:label ?cityName .
  FILTER (LANG(?cityName) = "en")
}
ORDER BY ?cityName""")

p = doc.add_paragraph()
r = p.add_run('Fourth query (Wikidata city attributes, equivalent of Task 2.2): ')
r.bold = True
r.font.color.rgb = BLACK
r2 = p.add_run(
    'Extends the third query with population (P1082), area (P2046) '
    'and geographic coordinates (P625).'
)
r2.font.color.rgb = BLACK

add_code_block(doc, """\
PREFIX wd:   <http://www.wikidata.org/entity/>
PREFIX wdt:  <http://www.wikidata.org/prop/direct/>
PREFIX rdfs: <http://www.w3.org/2000/01/rdf-schema#>

SELECT DISTINCT ?city ?cityName ?population ?area ?coord
WHERE {
  ?city wdt:P31/wdt:P279* wd:Q515 ;
        wdt:P17 wd:Q145 ;
        rdfs:label ?cityName .
  FILTER (LANG(?cityName) = "en")
  OPTIONAL { ?city wdt:P1082 ?population }
  OPTIONAL { ?city wdt:P2046 ?area }
  OPTIONAL { ?city wdt:P625  ?coord }
}
ORDER BY ?cityName""")

# ── save ──────────────────────────────────────────────────────────────────────
out_path = r'C:\Users\94540\Desktop\UoM\63502\cw2\CW2_Report.docx'
doc.save(out_path)
print(f'Saved to {out_path}')
