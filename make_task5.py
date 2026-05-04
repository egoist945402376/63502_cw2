from docx import Document
from docx.shared import Pt

doc = Document()

def add_normal(doc, text):
    return doc.add_paragraph(text)

doc.add_heading("Task 5 (Implementation)", level=3)

# ── Task 5.1 ───────────────────────────────────────────────────────────────
doc.add_heading("Task 5.1", level=4)

add_normal(doc, "(1) Two runs of Alg 1 with k = 5:")

table = doc.add_table(rows=3, cols=4)
table.style = "Table Grid"

for i, h in enumerate(["", "Run 1", "Run 2", "Average"]):
    run = table.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True

for i, v in enumerate(["Running Time (min)", "5.87", "6.67", "6.27"]):
    table.rows[1].cells[i].text = v
for i, v in enumerate(["Accuracy", "0.30", "0.35", "0.33"]):
    table.rows[2].cells[i].text = v

doc.add_paragraph()

add_normal(doc,
    "(2) Analysis of two columns with inconsistent annotations (≤60 words):\n\n"
    "CTRL_WIKI_GEO_member_states_united_nations col=0 (GT: dbo:Country, predicted: "
    "dbo:Place): Step 4 fails — DBpedia country entities declare both dbo:Country and "
    "superclass dbo:Place as rdf:type triples; dbo:Place accumulates higher frequency "
    "across sampled entities, overriding the correct fine-grained class.\n\n"
    "TOUGH_WEB_MISSP_drugs col=0 (GT: dbo:Drug, predicted: dbo:ChemicalSubstance): "
    "Step 2 fails — misspelled inputs (e.g., “acetomenophen”) cause the Lookup API "
    "to retrieve wrong entities; Step 4 then selects the superclass "
    "dbo:ChemicalSubstance over dbo:Drug."
)

# ── Task 5.2 ───────────────────────────────────────────────────────────────
doc.add_heading("Task 5.2", level=4)

add_normal(doc,
    "Weakness 1 – Superclass dominance: DBpedia entities declare multiple rdf:type "
    "triples; coarse superclasses (dbo:Place, dbo:Agent) accumulate higher frequency "
    "in Step 4, overriding fine-grained correct types like dbo:Country.\n\n"
    "Weakness 2 – No noise tolerance: Step 2 returns one entity per cell with no "
    "fallback; misspelled or noisy values retrieve wrong entities, corrupting "
    "Steps 3 and 4 with no recovery."
)

out = "C:/Users/94540/Desktop/UoM/63502/cw2/Task5_Report.docx"
doc.save(out)
print("Saved:", out)
